#!/usr/bin/env python3

import argparse
import html
import re
import shutil
import subprocess
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "1F4E79"
DARK_BLUE = "17365D"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "1F2937"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

DOC_META = {
    "01-role-model": ("Ролевая модель", "Матрица доступа, ответственности и разделения полномочий"),
    "02-security": ("Безопасность платформы", "Модель угроз, защитные меры и требования к промышленной эксплуатации"),
    "03-asutp-isolated-flow-model": ("Изолированный контур АСУ ТП", "Модель потоков данных и правила межсегментного взаимодействия"),
    "04-external-network-interactions": ("Внешнее и сетевое взаимодействие", "Интеграции, протоколы, порты и границы доверия"),
    "05-licensing": ("Лицензирование", "Лицензионная модель компонентов и закупочные рекомендации"),
    "06-infrastructure": ("Инфраструктурная схема", "Развёртывание, ресурсы, отказоустойчивость и эксплуатация"),
    "07-conceptual-application": ("Концептуальная схема приложения", "Компоненты, доменная модель и сквозные процессы"),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    no_split.set(qn("w:val"), "true")
    tr_pr.append(no_split)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(color)
    run_pr.append(underline)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>|https?://\S+)")


def add_inline(paragraph, text, font_size=None, color=DARK):
    text = html.unescape(text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n"))
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            style_run(run, font_size, color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            style_run(run, font_size, color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            style_run(run, font_size, color, font="Liberation Mono")
            run.font.highlight_color = None
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("<"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        else:
            url = token.rstrip(".,;)")
            add_hyperlink(paragraph, url, url)
            tail = token[len(url):]
            if tail:
                run = paragraph.add_run(tail)
                style_run(run, font_size, color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        style_run(run, font_size, color)


def style_run(run, size=None, color=DARK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def configure_document(doc, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 6, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("ДИСПЕТЧЕР  /  ТЕХНИЧЕСКАЯ ДОКУМЕНТАЦИЯ")
    style_run(run, 8.5, MID_GRAY, bold=True)
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    run = footer_paragraph.add_run(title + "  •  ")
    style_run(run, 8.5, MID_GRAY)
    add_page_field(footer_paragraph)


def parse_markdown(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("```"):
            language = line[3:].strip()
            payload = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                payload.append(lines[index])
                index += 1
            index += 1
            blocks.append(("code", language, "\n".join(payload)))
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            blocks.append(("heading", len(heading.group(1)), heading.group(2).strip()))
            index += 1
            continue
        if line.startswith(">"):
            payload = []
            while index < len(lines) and lines[index].startswith(">"):
                payload.append(lines[index].lstrip("> "))
                index += 1
            blocks.append(("quote", " ".join(payload)))
            continue
        if re.match(r"^\s*[-*+]\s+", line):
            items = []
            while index < len(lines) and re.match(r"^\s*[-*+]\s+", lines[index]):
                items.append(re.sub(r"^\s*[-*+]\s+", "", lines[index]).strip())
                index += 1
            blocks.append(("list", "bullet", items))
            continue
        if re.match(r"^\s*\d+[.)]\s+", line):
            items = []
            while index < len(lines) and re.match(r"^\s*\d+[.)]\s+", lines[index]):
                items.append(re.sub(r"^\s*\d+[.)]\s+", "", lines[index]).strip())
                index += 1
            blocks.append(("list", "number", items))
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[index + 1]):
            rows = []
            header_row = split_table_row(line)
            index += 2
            rows.append(header_row)
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append(("table", rows))
            continue
        if re.match(r"^(-{3,}|\*{3,})$", line.strip()):
            blocks.append(("rule",))
            index += 1
            continue
        paragraph = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if not candidate.strip():
                break
            if candidate.startswith(("#", "```", ">", "|")) or re.match(r"^\s*([-*+]\s+|\d+[.)]\s+)", candidate):
                break
            paragraph.append(candidate.strip())
            index += 1
        blocks.append(("paragraph", " ".join(paragraph)))
    return blocks


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_cover(doc, title, subtitle, source_name):
    for _ in range(3):
        doc.add_paragraph().paragraph_format.space_after = Pt(18)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("ПЛАТФОРМА УПРАВЛЕНИЯ ОПОВЕЩЕНИЯМИ «ДИСПЕТЧЕР»")
    style_run(run, 10, BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(20)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    style_run(run, 28, DARK_BLUE, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(44)
    run = paragraph.add_run(subtitle)
    style_run(run, 14, MID_GRAY)
    meta = doc.add_table(rows=4, cols=2)
    widths = [2300, 7060]
    values = [
        ("Статус", "Архитектурно-техническая документация"),
        ("Основание", "Фактическая реализация репозитория и целевые требования production-контура"),
        ("Версия", "1.0 • 10 августа 2026 г."),
        ("Исходный файл", source_name),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = ""
        add_inline(row.cells[0].paragraphs[0], label, 9, MID_GRAY)
        row.cells[1].text = ""
        add_inline(row.cells[1].paragraphs[0], value, 9.5, DARK)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    set_table_geometry(meta, widths)
    doc.add_page_break()


def add_contents(doc, blocks):
    paragraph = doc.add_paragraph("Содержание", style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(0)
    entries = [block for block in blocks if block[0] == "heading" and block[1] != 1]
    compact = len(entries) > 28
    seen = set()
    for block in blocks:
        if block[0] != "heading" or block[1] == 1:
            continue
        label = re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", block[2]).strip()
        key = (block[1], label)
        if key in seen:
            continue
        seen.add(key)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18 * max(0, block[1] - 2))
        p.paragraph_format.space_after = Pt(1.5 if compact else 3)
        run = p.add_run(label)
        if compact:
            font_size = 9 if block[1] == 2 else 8.5
        else:
            font_size = 10 if block[1] == 2 else 9.5
        style_run(run, font_size, DARK, bold=block[1] == 2)
    doc.add_page_break()


def status_fill(text):
    normalized = text.upper()
    if "ФАКТ" in normalized or "РЕАЛИЗОВАН" in normalized:
        return GREEN
    if "ЦЕЛЕВО" in normalized or "РЕКОМЕНД" in normalized:
        return LIGHT_BLUE
    if "РАЗРЫВ" in normalized or "РИСК" in normalized or "НЕ РЕАЛИЗ" in normalized:
        return RED
    if "СОГЛАСОВ" in normalized or "ПЛАН" in normalized:
        return AMBER
    return None


def table_widths(rows):
    columns = max(len(row) for row in rows)
    scores = []
    for column in range(columns):
        values = [row[column] if column < len(row) else "" for row in rows]
        length = max(6, min(45, max(len(re.sub(r"[*`]", "", value)) for value in values)))
        scores.append(length)
    total = sum(scores)
    widths = [max(650, round(TABLE_WIDTH_DXA * score / total)) for score in scores]
    difference = TABLE_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 650:
        deficit = 650 - widths[-1]
        widths[-1] = 650
        widest = max(range(len(widths) - 1), key=lambda index: widths[index])
        widths[widest] -= deficit
    return widths


def add_table(doc, rows):
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    font_size = 8.5 if columns <= 5 else 7.25
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        for column_index in range(columns):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1.5)
            paragraph.paragraph_format.line_spacing = 1.05
            value = values[column_index] if column_index < len(values) else ""
            add_inline(paragraph, value, font_size, WHITE if row_index == 0 else DARK)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            fill = status_fill(value)
            if row_index > 0 and fill:
                set_cell_shading(cell, fill)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, table_widths(rows))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.line_spacing = Pt(1)


def add_code(doc, language, payload):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(payload)
    style_run(run, 8.5, DARK, font="Liberation Mono")
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("Листинг" + (" • " + language if language else ""))
    style_run(run, 8.5, MID_GRAY, italic=True)


def mermaid_to_dot(payload):
    direction = "LR" if re.search(r"(?:flowchart|graph)\s+LR", payload) else "TB"
    nodes = {}
    edges = []
    node_pattern = re.compile(r'\b([A-Za-zА-Яа-я_][\w-]*)\s*(?:\[\(|\[|\{|\()\s*["\']?(.+?)["\']?\s*(?:\)\]|\]|\}|\))(?=\s|$)')
    for line in payload.splitlines():
        cleaned = line.strip()
        if not cleaned or cleaned.startswith(("flowchart ", "graph ", "subgraph", "end", "style", "classDef")):
            continue
        for match in node_pattern.finditer(cleaned):
            label = match.group(2).strip().strip('"\'').replace("\\n", "\n").replace("<br/>", "\n").replace("<br>", "\n")
            if "-->" in label or "-.->" in label:
                continue
            nodes[match.group(1)] = label
        edge_match = re.search(r'([\w-]+)(?:\s*(?:\[[^\]]+\]|\([^)]*\)|\{[^}]*\}))?\s*(-\.->|-->|==>)\s*(?:\|([^|]+)\|\s*)?([\w-]+)', cleaned)
        if edge_match:
            edges.append((edge_match.group(1), edge_match.group(4), edge_match.group(3) or "", edge_match.group(2) == "-.->"))
            nodes.setdefault(edge_match.group(1), edge_match.group(1))
            nodes.setdefault(edge_match.group(4), edge_match.group(4))
    if not edges:
        return None
    if len(nodes) > 8:
        direction = "TB"
    lines = [
        "digraph G {",
        f"rankdir={direction};",
        'graph [bgcolor="white", pad="0.2", nodesep="0.35", ranksep="0.5"];',
        'node [shape=box, style="rounded,filled", fillcolor="#DCE6F1", color="#1F4E79", fontname="Arial", fontsize=10, margin="0.12,0.08"];',
        'edge [color="#667085", fontcolor="#475467", fontname="Arial", fontsize=8, arrowsize=0.7];',
    ]
    for identifier, label in nodes.items():
        safe = label.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n")
        lines.append(f'"{identifier}" [label="{safe}"];')
    for source, target, label, dashed in edges:
        attrs = []
        if label:
            attrs.append('label="' + label.replace('"', '\\"') + '"')
        if dashed:
            attrs.append('style="dashed"')
        suffix = " [" + ", ".join(attrs) + "]" if attrs else ""
        lines.append(f'"{source}" -> "{target}"{suffix};')
    lines.append("}")
    return "\n".join(lines)


def add_mermaid(doc, payload, asset_dir, index):
    dot_source = mermaid_to_dot(payload)
    if dot_source is None or shutil.which("dot") is None:
        add_code(doc, "mermaid", payload)
        return
    dot_path = asset_dir / f"diagram-{index:02d}.dot"
    png_path = asset_dir / f"diagram-{index:02d}.png"
    dot_path.write_text(dot_source, encoding="utf-8")
    result = subprocess.run(["dot", "-Tpng", "-Gdpi=170", str(dot_path), "-o", str(png_path)], capture_output=True, text=True)
    if result.returncode != 0 or not png_path.exists():
        add_code(doc, "mermaid", payload)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    with Image.open(png_path) as image:
        aspect = image.width / image.height
    if aspect < 6.25 / 6.5:
        shape = run.add_picture(str(png_path), height=Inches(6.5))
    else:
        shape = run.add_picture(str(png_path), width=Inches(6.25))
    shape._inline.docPr.set("title", "Архитектурная схема %s" % index)
    shape._inline.docPr.set("descr", "Схема, автоматически построенная по проверенному Markdown-описанию")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(f"Схема {index}")
    style_run(run, 8.5, MID_GRAY, italic=True)


def render_blocks(doc, blocks, asset_dir):
    diagram_index = 0
    first_h1_skipped = False
    current_heading = ""
    for block in blocks:
        kind = block[0]
        if kind == "heading":
            level, text = block[1], block[2]
            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
                continue
            current_heading = text.lower()
            style = "Heading " + str(min(4, max(1, level - 1)))
            paragraph = doc.add_paragraph(style=style)
            add_inline(paragraph, text, None, BLUE if level <= 2 else DARK_BLUE)
        elif kind == "paragraph":
            paragraph = doc.add_paragraph()
            if "вывод" in current_heading or "заключение" in current_heading:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(4)
                add_inline(paragraph, block[1], 10)
            else:
                add_inline(paragraph, block[1])
        elif kind == "quote":
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, LIGHT_BLUE)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, block[1], 10.5, DARK_BLUE)
            set_table_geometry(table, [TABLE_WIDTH_DXA])
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(1)
        elif kind == "list":
            style = "List Bullet" if block[1] == "bullet" else "List Number"
            for item in block[2]:
                paragraph = doc.add_paragraph(style=style)
                add_inline(paragraph, item)
        elif kind == "table":
            add_table(doc, block[1])
        elif kind == "code":
            if block[1].lower() == "mermaid":
                diagram_index += 1
                add_mermaid(doc, block[2], asset_dir, diagram_index)
            else:
                add_code(doc, block[1], block[2])
        elif kind == "rule":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph_border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "B4C7E7")
            paragraph_border.append(bottom)
            paragraph._p.get_or_add_pPr().append(paragraph_border)


def build_document(source, output, asset_root):
    key = source.stem
    title, subtitle = DOC_META.get(key, (key, "Техническая документация"))
    blocks = parse_markdown(source)
    doc = Document()
    configure_document(doc, title)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "Проект «Диспетчер»"
    doc.core_properties.comments = "Сформировано по фактической реализации репозитория и архитектурным требованиям."
    add_cover(doc, title, subtitle, source.name)
    add_contents(doc, blocks)
    asset_dir = asset_root / key
    asset_dir.mkdir(parents=True, exist_ok=True)
    render_blocks(doc, blocks, asset_dir)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description="Сборка комплекта архитектурной документации в DOCX")
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--asset-dir", type=Path, required=True)
    args = parser.parse_args()
    missing = []
    for key in DOC_META:
        source = args.source_dir / (key + ".md")
        if not source.exists():
            missing.append(str(source))
            continue
        build_document(source, args.output_dir / (key + ".docx"), args.asset_dir)
    if missing:
        raise SystemExit("Не найдены исходники:\n" + "\n".join(missing))


if __name__ == "__main__":
    main()
